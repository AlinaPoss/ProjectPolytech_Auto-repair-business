from tkinter import *
from tkinter import  messagebox
import random
import docx

personal = {}
posetit = {}
sklad={}
cars = {}
service = {}
orders = {}
masters = {}
bool_reg = 0

def get_remind():
    global firma
    with open("remind.txt", "r", encoding="utf-8") as file:
        lines = file.readlines()
    firma = lines[0]
    global report_remind_posetit
    global report_remind_sklad
    report_remind_sklad = lines[1]
    report_remind_posetit  = lines[2]

def get_orders():
    global id_ord
    id_ord = []
    global id_ord_chel
    id_ord_chel = []
    global id_ord_delo
    id_ord_delo = []
    global id_ord_mast
    id_ord_mast = []
    global id_ord_car
    id_ord_car = []
    with open("orders.txt", "r", encoding="utf-8") as file:
         for line in file:
            content = line.split('///')
            li = [content[i] for i in range(1, 5)]
            orders[content[0]] = li
            id_ord.append(content[0])
            id_ord_chel.append(content[1])
            id_ord_delo.append(content[2])
            id_ord_mast.append(content[3])
            id_ord_car.append(content[4])

def get_masters():
    global id_master
    id_master = []
    global name_master
    name_master = []
    global secname_master
    secname_master = []
    with open("masters.txt", "r", encoding="utf-8") as file:
        for line in file:
            content = line.split('///')
            li = [content[i] for i in range(1, 3)]
            masters[content[0]] = li
            id_master.append(content[0])
            name_master.append(content[1])
            secname_master.append(content[2])

def get_cars():
    global car_id
    car_id = []
    global car_id_model
    car_id_model = []
    global car_gos_numb
    car_gos_numb = []
    global car_vin
    car_vin = []
    global car_id_pers
    car_id_pers = []
    with open("cars.txt", "r", encoding="utf-8") as file:
        for line in file:
            content = line.split('///')
            li = [content[i] for i in range(1, 5)]
            cars[content[0]] = li#list
            car_id.append(content[0])
            car_id_model.append(content[1])
            car_gos_numb.append(content[2])
            car_vin.append(content[3])
            car_id_pers.append(content[4])

def get_service():
    global id_serv
    id_serv = []
    global work
    work = []
    global work_price
    work_price = []
    global need_detals
    need_detals = []
    with open("service.txt", "r", encoding="utf-8") as file:
        for line in file:
            content = line.split('///')
            li = [content[i] for i in range(1, 3)]
            id_serv.append(content[0])
            work.append(content[1])
            work_price.append(content[2])
            for i in range(3, len(content)):
                need_detals.append(content[i])
            li.append(need_detals)
            service[content[0]] = li

def get_sklad():
    global attr_id
    attr_id = []
    global attribute
    attribute = []
    global kol
    kol = []
    global attr_id_model
    attr_id_model = []
    global attr_price
    attr_price = []
    with open("sklad.txt", "r", encoding="utf-8") as file:
        for line in file:
            content = line.split('///')
            li = [content[i] for i in range(1, 5)]
            sklad[content[0]] = li#list
            attr_id.append(content[0])
            attribute.append(content[1])
            kol.append(content[2])
            attr_id_model.append(content[3])
            attr_price.append(content[4])

        
def get_people():
    global i_d
    i_d = []
    global log
    log = []
    global pas
    pas = []
    global name
    name = []
    global secname
    secname = []
    global post
    post = []
    with open("personal.txt", "r", encoding="utf-8") as file:
        for line in file:
            content = line.split('///')
            li = [content[i] for i in range(1, 6)]
            personal[content[0]]= li#list
            i_d.append(content[0])
            log.append(content[1])
            pas.append(content[2])
            name.append(content[3])
            secname.append(content[4])
            post.append(content[5])

    global id_poset
    id_poset = []
    global log_poset
    log_poset = []
    global pas_poset
    pas_poset = []
    global name_poset
    name_poset = []
    global secname_poset
    secname_poset = []
    global mobile_poset
    mobile_poset = []
    global poset_id_car
    poset_id_car = []
    with open("posetit.txt", "r", encoding="utf-8") as file:
        for line in file:
            content = line.split('///')
            li = [content[i] for i in range(1, 7)]#list
            posetit[content[0]] = li
            id_poset.append(content[0])
            log_poset.append(content[1])
            pas_poset.append(content[2])
            name_poset.append(content[3])
            secname_poset.append(content[4])
            poset_id_car.append(content[5])
            mobile_poset.append(content[6])

def da_firma_func():
    fi = e3.get()
    firma = fi
    e3.delete(0, END)
    title_admin.delete('all')
    title_admin.create_text(200,10, text=firma, fill="slateblue")


def net_firma_func():
    e3.delete(0, END)
    
def quest_firma():
    res = messagebox.askquestion("!", "Вы уверены, что хотите поменять название?", icon='question' )
    if res == 'yes':
        da_firma_func()
    else:
        net_firma_func()

def new_pass():
    log_to_change = e8.get()
    pass_to_change = e9.get()
    e8.pack_forget()
    e9.pack_forget()
    but_pass_q.pack_forget()
    if log_to_change in log:
        res = messagebox.askquestion("!", "Вы уверены, что хотите поменять пароль этого пользователя?", icon='question' )
        if res == 'yes':
            ind = log.index(log_to_change)
            pas[ind]=pass_to_change
            key = i_d[ind]
            personal[key][1] = str(pass_to_change)
            string = str(key)+'///'
            for pos in personal[key]:
                string = string + str(pos)+ '///'
            with open('personal.txt', 'r', encoding = 'utf-8') as file:
                lines = file.readlines()
            lines[ind] = string
            with open('personal.txt', 'w', encoding = 'utf-8') as file:
                for pos in lines:
                    file.write(pos)
            messagebox.showinfo(title='Результат', message='Пароль успешно изменён')
            
    elif log_to_change in log_poset:
        res = messagebox.askquestion("!", "Вы уверены, что хотите поменять пароль этого пользователя?", icon='question' )
        if res == 'yes':
            ind = log_poset.index(log_to_change)
            pas_poset[ind]=pass_to_change
            key = id_poset[ind]
            posetit[key][1] = str(pass_to_change)
            string = str(key)+'///'
            for pos in posetit[key]:
                string = string + str(pos)+ '///'
            with open('posetit.txt', 'r', encoding = 'utf-8') as file:
                lines = file.readlines()
            lines[ind] = string
            with open('posetit.txt', 'w', encoding = 'utf-8') as file:
                for pos in lines:
                    file.write(pos)
            messagebox.showinfo(title='Результат', message='Пароль успешно изменён')
    else:
         messagebox.showerror('Ошибка','Пользователь не найден')


def change_pass_func():
    global e8
    e8 = Entry(change_pass_frame, width=50)
    e8.insert(0, 'Введите логин пользователя')
    e8.pack()
    global e9
    e9 = Entry(change_pass_frame, width=50)
    e9.insert(0, 'Новый пароль')
    e9.pack()

    global but_pass_q
    but_pass_q = Button(change_pass_frame, text='Поменять', command=new_pass)
    but_pass_q.pack()

def enter_admin():
    global okno_admin
    okno_admin = Frame(root, width = 600, height=600,bg = 'peachpuff')
    okno_admin.pack()
    global title_admin
    title_admin = Canvas(okno_admin, width=600, height=30)
    title_admin.pack()
    title_admin.create_text(300,35, text=firma,  fill="slateblue", justify=CENTER, font = 'Times 20')
    global e3
    e3 = Entry(okno_admin, width=50)
    e3.pack()
    e3.insert(0, 'Введите новое название')
    change_but = Button(okno_admin, text='Поменять название автосервиса', command=quest_firma)
    change_but.pack()
    title_2 = Canvas(okno_admin, width=600, height=30)
    title_2.pack()
    global change_pass
    change_pass = Button(okno_admin, text='Поменять пароль пользователя', command=change_pass_func)
    change_pass.pack()
    global change_pass_frame
    change_pass_frame = Frame(okno_admin, width = 600, height=40,bg = 'peachpuff')
    change_pass_frame.pack()

def detals_in_stock(id_delo):
    li = service[id_delo][2]
    for i in li:
        if int(sklad[i][2])<1:
            return 0
    return 1

def create_new_order():###########
    get_masters()
    get_service()
    get_sklad()
    get_orders()
    delo = e_service.get() 
    if delo in work:
        id_order = random.randint(10, 10000)+random.randint(10, 100000)
        id_mast = id_master[0]
        id_chel = id_poset[log_poset.index(s_log)]
        id_car = poset_id_car[log_poset.index(s_log)]
        for i in id_serv:
            if service[i][0]==delo:
                id_delo = id_serv[work.index(delo)]
        if detals_in_stock(id_delo):
            li = []
            id_ord.append(id_order)
            id_ord_chel.append(id_chel)
            id_ord_delo.append(id_order)
            id_ord_mast.append(id_mast)
            id_ord_car.append(id_car)
            li.append(id_chel)
            li.append(id_delo)
            li.append(id_mast)
            li.append(id_car)
            orders[id_order] = li
            string = str(id_order)
            for i in li:
                string = string+str(i)+'///'
            with open("orders.txt", "a", encoding="utf-8") as file:
                file.write(string+'\n')
            for i in service[id_delo][2]:
                sklad[i][1] = int(sklad[i][1])-1
                ind = attr_id.index(i)
                kol[ind] = int(kol[ind])-1
            messagebox.showinfo('Результат','Заказ совершен успешно')
        else:
            messagebox.showinfo('Оповещение','К сожалению, необхожимые детали отсутствуют на складе')
        
    else:
        messagebox.showerror('Ошибка','Услуга не найдена')
    create_but.pack_forget()
    e_service.pack_forget()
        
def func_new_order():
    global e_service
    e_service = Entry(frame_new_order, width=50)
    e_service.insert(0,'Укажите вид работы')
    e_service.pack()
    global create_but
    create_but = Button(frame_new_order, text='Оформить заказ', command=create_new_order) 
    create_but.pack()

def func_show_orders():
    get_masters()
    get_service()
    get_sklad()
    get_cars()
    get_orders()
    doc = docx.Document()
    for i in orders:
        if orders[i][0] == id_poset[log_poset.index(s_log)]:
            key_ord = i
    string = 'Заказчик: '+str(name_poset[log_poset.index(s_log)])+' '+str(secname_poset[log_poset.index(s_log)])
    doc.add_paragraph(string)
    string = 'Наименование работы: '+str(service[orders[key_ord][1]][0])
    doc.add_paragraph(string)
    string = 'Исполнитель: OOO "'+str(firma)+'"'
    doc.add_paragraph(string)
    #string = 'Модель машины : '+str(cars[poset_id_car[log_poset.index(s_log)]][0])
    #doc.add_paragraph(string)
    string = 'Необходимые детали: '
    for i in service[orders[key_ord][1]][2]:
        string = string+str(sklad[i][0])+', '
    string = string[0:-2]
    doc.add_paragraph(string)
    st = 0
    for i in service[orders[key_ord][1]][2]:
        st = st+int(sklad[i][3])
    string = 'Общая стоимость: '+str(st)+' руб.'
    doc.add_paragraph(string)
    
    doc_name = 'Заказ-наряды_'+str(secname_poset[log_poset.index(s_log)])+'.docx'
    doc.save(doc_name)
    messagebox.showinfo(title='Результат', message='Успешно скачан файл с вашими заказ-нарядами')

def enter_posetit():
    global okno_poset
    okno_poset = Frame(root, width = 600, height=600, bg = 'skyblue')
    okno_poset.pack(side=TOP)
    title = Canvas(okno_poset, width=600, height=30)
    title.pack(side=TOP)
    title.create_text(300,35, text=firma,  fill="slateblue", justify=CENTER, font="Times 20")

    global frame_new_order
    frame_new_order = Frame(okno_poset, width = 600, height=600, bg = 'skyblue')
    frame_new_order.pack()
    but_new_order = Button(frame_new_order, text='Создать новый заказ', command=func_new_order)
    but_new_order.pack(side=TOP)
    
    but_show_orders = Button(okno_poset, text='Показать мои заказы', command=func_show_orders)
    but_show_orders.pack(side=TOP)

def proverka():
    detal = e4.get()
    if detal in attribute:
        ind = attribute.index(detal)
        che_vishlo = 'Осталось '+str(kol[ind])+'шт.'
    else:
        che_vishlo = 'Деталь не найдена'
    result_check.delete('all')
    result_check.create_text(100, 20, text=che_vishlo, fill="red")

def change_sklad_deep():
    new_kol = e6.get()
    new_kol = '0'+new_kol[1:]
    min_kol = e7.get()
    min_kol = '0'+min_kol[1:]
    e6.pack_forget()
    e7.pack_forget()
    but_change_sklad.pack_forget()
    if new_kol.isdigit():
        ind = attribute.index(detal_change)
        kol[ind] = int(kol[ind])+int(new_kol)-int(min_kol)
        key = attr_id[ind]
        sklad[key][1] = str(kol[ind])
        string = str(key)+'///'
        for pos in sklad[key]:
            string = string + str(pos) + '///'
        with open("sklad.txt", "r", encoding="utf-8") as file:
            lines = file.readlines()
        lines[ind] = string
        with open("sklad.txt", "w", encoding="utf-8") as file:
            for pos in lines:
                file.write(pos)
        messagebox.showinfo(title='Результат', message='Количество деталей успешно изменено')
    else:
        messagebox.showinfo(title='Результат', message='Некорректные данные')

def change_sklad():
    global detal_change
    detal_change = e5.get()
    if detal_change in attribute:
        global e6
        e6 = Entry(check_fr, width=50)
        e6.pack(side=TOP)
        e6.insert(0, '+')
        global e7
        e7 = Entry(check_fr, width=50)
        e7.pack(side=TOP)
        e7.insert(0, '-')
        global but_change_sklad
        but_change_sklad = Button(check_fr, text='Изменить', command=change_sklad_deep)
        but_change_sklad.pack(side=TOP)
    else:
        messagebox.showinfo(title='Результат', message='Деталь не обнаружена на складе')

def make_sklad_report():
    global report_remind_sklad
    report_remind_sklad = int(report_remind_sklad)+1
    report_name = 'Отчет о состоянии склада_'+str(report_remind_sklad)+'.docx'
    doc = docx.Document()
    doc.add_paragraph('Деталь    '+'Количество ' + 'Цена(руб.) '+ '\n')
    for i in sklad:
        line = sklad[i]
        string = str(line[0])+'          '+str(line[1])+'                '+str(line[3])
        doc.add_paragraph(string)
    doc.save(report_name)
    messagebox.showinfo(title='Результат', message='Отчет успешно создан')
    
def enter_sklad():
    get_sklad()
    global okno_sklad
    okno_sklad = Frame(root, width = 600, height=600,bg = 'palegreen')
    okno_sklad.pack()
    title = Canvas(okno_sklad, width=600, height=30)
    title.pack(side=TOP)
    title.create_text(300,30, text=firma,  fill="slateblue", justify=CENTER, font="Times 20")
    global e4
    e4 = Entry(okno_sklad, width=50)
    e4.pack(side=TOP)
    e4.insert(0, 'Название атрибута')
    global check_is
    check_is = Button(okno_sklad, text='Проверить наличие на складе', command=proverka)
    check_is.pack(side=TOP)
    global result_check
    result_check = Canvas(okno_sklad, width=600, height=30)
    result_check.pack(side=TOP)
    result_check.create_text(100, 20, text='', fill="red")
    global e5
    e5 = Entry(okno_sklad, width=50)
    e5.pack(side=TOP)
    e5.insert(0, 'Название атрибута')
    global change_is
    change_is = Button(okno_sklad, text='Изменить количество деталей', command=change_sklad)
    change_is.pack(side=TOP)
    
    global check_fr
    check_fr = Frame(okno_sklad, width=600, height=60, bg = 'lemonchiffon')
    check_fr.pack(side=TOP)
    
    sklad_report = Button(okno_sklad, text='Создать отчет о состоянии склада', command=make_sklad_report)
    sklad_report.pack(side=TOP)

def func_know_price():
    detal = e_know_price.get()
    if detal in attribute:
        ind = attribute.index(detal)
        che_vishlo = 'Цена детали: '+str(attr_price[ind])+'руб.'
    else:
        che_vishlo = 'Деталь не найдена'
    kol_det.delete('all')
    kol_det.create_text(100, 20, text=che_vishlo, fill="black", font = 'Times 15')

def enter_account():
    get_sklad()
    global okno_account
    okno_account = Frame(root,  width = 1800, height=1600, bg = 'thistle')
    okno_account.pack()
    title = Canvas(okno_account, width=1800, height=40)
    title.create_text(300, 35, text=firma,  fill="slateblue", justify=CENTER, font = 'Times 20')
    title.pack(side=TOP)
    
    sklad_report = Button(okno_account, text='Создать отчет о состоянии склада', command=make_sklad_report)
    sklad_report.pack(side=TOP)
    cheta = Canvas(okno_account, width=1800, height=40, bg = 'thistle')
    cheta.pack()
    
    fr_know_price = Frame(okno_account, width = 1800, height=600,bg = 'thistle')
    fr_know_price.pack()
    global e_know_price
    e_know_price = Entry(fr_know_price, width=50)
    e_know_price.insert(0, 'Название детали')
    e_know_price.pack()
    but_know = Button(fr_know_price, text='Узнать цену детали', command=func_know_price)
    but_know.pack()
    global kol_det
    kol_det = Canvas(fr_know_price, width=1800, height=30, bg = 'thistle')
    kol_det.pack(side=TOP)
    kol_det.create_text(100, 20, text='', fill="black")

def func_know_poset():
    text_know_poset = str(len(posetit))+' чел.'
    tit_kol.create_text(300, 20, text=text_know_poset,  fill="slateblue", justify=CENTER, font = 'Times 20')

def func_make_order():
    global report_remind_posetit
    get_orders()
    report_remind_posetit = int(report_remind_posetit)+1
    report_name = 'Отчет о посетителях_'+str(report_remind_posetit)+'.docx'
    with open("remind.txt", "r", encoding="utf-8") as file:
        lines = file.readlines()
    lines[2] = str(report_remind_posetit+1)
    with open("remind.txt", "w", encoding="utf-8") as file:
        for i in lines:
             file.write(i)
    doc = docx.Document()
    doc.add_paragraph('Имя      Фамилия      Количество заказов   '+ '\n')
    ko=0
    for i in posetit:
        string = ''
        ko = 0
        for k in orders:
            if orders[k][0]==i:
                ko=ko+1
        koo = str(ko)
        string = string +str(posetit[i][2])+'  '+str(posetit[i][3])+'   '+koo
        doc.add_paragraph(string)
    doc.save(report_name)
    messagebox.showinfo(title='Результат', message='Отчет успешно создан')
    

def enter_meneg():
    global okno_meneg
    okno_meneg = Frame(root,  width = 1800, height=1600, bg = 'thistle')
    okno_meneg.pack()
    title = Canvas(okno_meneg, width=1800, height=40)
    title.create_text(300,35, text=firma,  fill="slateblue", justify=CENTER, font = 'Times 20')
    title.pack(side=TOP)
    but_know = Button(okno_meneg, text='Узнать количетво пользователей', command=func_know_poset)
    but_know.pack()
    global tit_kol
    tit_kol = Canvas(okno_meneg, width=1800, height=40, bg = 'thistle')
    tit_kol.pack()
    but_make_order = Button(okno_meneg, text='Создать отчет о пользователях', command=func_make_order)
    but_make_order.pack()
    
def insert():
    global s_log
    if bool_reg:
        s_log = e_log.get()
        s_pass = e_pas.get()
    else:
        s_log = e1.get()
        s_pass = e2.get()
    if s_log in log:
        if s_pass == pas[log.index(s_log)]:
            
            okno_enter.pack_forget()
            if post[log.index(s_log)] == 'admin':
                enter_admin()
            elif post[log.index(s_log)] == 'sklad':
                enter_sklad()
            elif post[log.index(s_log)] == 'account':
                enter_account()
            elif post[log.index(s_log)] == 'meneg':
                enter_meneg()
        else:
            messagebox.showinfo(title='Предупреждение', message='Неверный пароль')
    elif s_log in log_poset:
         if s_pass == pas_poset[log_poset.index(s_log)]:
            okno_enter.pack_forget()
            enter_posetit()
         else:
            messagebox.showinfo(title='Предупреждение', message='Неверный пароль')
    else:
        messagebox.showinfo(title='Предупреждение', message='Пользователь не найден')

def reg_new_func():
    if e_name.get() in log_poset:
        messagebox.showerror('Ошибка','Такой пользователь уже существует')
    elif not e_name.get().isalpha():
        messagebox.showerror('Ошибка','Некорректное имя')
    elif not e_secname.get().isalpha():
        messagebox.showerror('Ошибка','Некоррекнтая фамилия')
    elif not e_mobile.get().isdigit():
        messagebox.showerror('Ошибка','Некорректный номер телефона')
    elif not (e_vin.get().isdigit() and len(str(e_vin.get()))==17):
        messagebox.showerror('Ошибка','Некорректный VIN')
    elif not e_gos_reg_numb.get().isdigit():
        messagebox.showerror('Ошибка','Неккоректный гос.номер')
    else:
        new_poset.pack_forget()
        id_chel = random.randint(10, 1000)+random.randint(10, 1000)
        id_car = random.randint(10, 1000)+random.randint(10, 1000)
        li = []
        li.append(e_model.get())
        li.append(e_gos_reg_numb.get())
        li.append(e_vin.get())
        li.append(id_chel)
        cars[id_car] = li
        string = str()+'///'+str(e_model.get())+'///'
        string= string+str(e_gos_reg_numb.get())+'///'+str(e_vin.get())+'///'+str(id_chel)+'///'
        with open('cars.txt', 'a', encoding = 'utf-8') as file:
            file.write('\n')
            file.write(string)

        li = []
        li.append(e_log.get())
        li.append(e_pas.get())
        li.append(e_name.get())
        li.append(e_secname.get())
        li.append(e_mobile.get())
        li.append(id_car)
        posetit[id_chel] = li
        string = str(id_chel)+'///'+str(e_log.get())+'///'+str(e_pas.get())+'///'
        string = string + str(e_name.get())+'///'+str(e_secname.get())+'///'+str(e_mobile.get())+'///'+str(id_car)+'///'
        with open('posetit.txt', 'a', encoding = 'utf-8') as file:
            file.write('\n')
            file.write(string)
        
        e_log.pack_forget()
        e_pas.pack_forget()
        e_name.pack_forget()
        e_secname.pack_forget()
        e_mobile.pack_forget()
        e_vin.pack_forget()
        e_gos_reg_numb.pack_forget()
        e_model.pack_forget()
        messagebox.showinfo(title='Результат', message='Пользователь успешно зарегистрирован')
        reg_frame.pack_forget()
        insert()

def registration():
    but_ins.grid_forget()
    e1.grid_forget()
    e2.grid_forget()
    reg_but.grid_forget()
    global bool_reg
    bool_reg = 1
    global reg_frame
    reg_frame = Frame(root, width = 50, height=60, bg='burlywood')
    reg_frame.pack()
    global e_log
    e_log = Entry(reg_frame, width=50)
    e_log.insert(0, 'Логин')
    e_log.pack()
    global e_pas
    e_pas = Entry(reg_frame, width=50)
    e_pas.insert(0, 'Пароль')
    e_pas.pack()
    global e_name
    e_name = Entry(reg_frame, width=50)
    e_name.insert(0, 'Имя')
    e_name.pack()
    global e_secname
    e_secname = Entry(reg_frame, width=50)
    e_secname.insert(0, 'Фамилия')
    e_secname.pack()
    global e_mobile
    e_mobile = Entry(reg_frame, width=50)
    e_mobile.insert(0, 'Номер телефона')
    e_mobile.pack()
    global e_vin
    e_vin = Entry(reg_frame, width=50)
    e_vin.insert(0, 'VIN машины')
    e_vin.pack()
    global e_gos_reg_numb
    e_gos_reg_numb = Entry(reg_frame,width=50)
    e_gos_reg_numb.insert(0, 'Государственный регистрационный номер машины')
    e_gos_reg_numb.pack()
    global e_model
    e_model = Entry(reg_frame, width=50)
    e_model.insert(0, 'Модель машины')
    e_model.pack()
    global new_poset
    new_poset = Button(reg_frame, text='Создать аккаунт', command=reg_new_func)
    new_poset.pack()

get_people()
get_remind()
global report_remind
report_remind = 0
root = Tk()
root.geometry('600x400')
okno_enter = Frame(root, width = 1800, height=1600, bg='powderblue')
global title
title = Canvas(okno_enter, width=700, height=50)
title.grid(column = 0, row=0, columnspan=3)
title.create_text(200, 35, text=firma,  fill="slateblue", justify=CENTER, font = 'Times 20')
e1 = Entry(okno_enter, width=50)  
e1.grid( column = 0, row=1)
e1.insert(0, 'Введите логин')
e2 = Entry(okno_enter, width=50)  
e2.grid(column=0, row=2)
e2.insert(0, 'Введите пароль')
global but_ins
but_ins = Button(okno_enter, text='Далее', command=insert)
but_ins.grid(column = 0, row=3)
okno_enter.pack()
global reg_but
reg_but = Button(okno_enter, text='Зарегистрироваться', command=registration)
reg_but.grid(column = 0, row=4)
root.mainloop()
